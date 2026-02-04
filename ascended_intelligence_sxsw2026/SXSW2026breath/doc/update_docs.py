#!/usr/bin/env python3
"""Update all docx files in doc/ with current content."""
from pathlib import Path
from docx import Document
from docx.shared import Pt

DOC_DIR = Path(__file__).resolve().parent


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p


def doc_readme():
    """Main README.docx"""
    doc = Document()
    doc.add_heading("Project: Combined Breath + Emotion Pipeline", 0)
    doc.add_paragraph(
        "This project runs a combined pipeline: Breath (OpenSMILE) and emotion2vec (Audio2Emotion) "
        "on the same audio, outputting breath state and emotion per segment. OSC output is sent to TouchDesigner."
    )
    doc.add_heading("Project layout", level=1)
    for line in [
        "ascended_intelligence_sxsw2026/ – Breath detector (OpenSMILE, F0, BPM)",
        "bridge/ – Combines Breath + emotion model; OSC output; prepare_audio",
        "english_model/ – emotion2vec, noise model, download/load",
        "test_recorded_audio.py – Test on recorded audio",
        "test_live_mic.py – Live microphone with VAD, enhancement, OSC",
        "osc_receiver.py – Test OSC without TouchDesigner",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("Install", level=1)
    doc.add_paragraph("pip install -r requirements.txt python-osc")
    doc.add_paragraph("pip install -e ./opensmile-python-main")
    doc.add_paragraph("conda install -y ffmpeg -c conda-forge")
    doc.add_heading("Run", level=1)
    doc.add_paragraph("Recorded: python test_recorded_audio.py path/to/audio.wav")
    doc.add_paragraph("Live mic: python test_live_mic.py")
    doc.add_paragraph("OSC test: python osc_receiver.py (terminal 1), python test_live_mic.py (terminal 2)")
    doc.add_heading("Documentation", level=1)
    doc.add_paragraph("bridge/README.md, english_model/README.md, TESTING.md, doc/TouchDesigner_Integration.docx")
    doc.save(DOC_DIR / "README.docx")


def doc_testing():
    """TESTING.docx"""
    doc = Document()
    doc.add_heading("Testing", 0)
    doc.add_heading("test_recorded_audio.py", level=1)
    doc.add_paragraph("python test_recorded_audio.py")
    doc.add_paragraph("python test_recorded_audio.py path/to/audio.wav")
    doc.add_heading("test_live_mic.py", level=1)
    doc.add_paragraph("python test_live_mic.py")
    doc.add_paragraph("python test_live_mic.py --osc-ip 127.0.0.1 --osc-port 5005")
    doc.add_paragraph("python test_live_mic.py --no-osc")
    doc.add_heading("osc_receiver.py", level=1)
    doc.add_paragraph("Run in terminal 1 while test_live_mic runs in terminal 2")
    doc.save(DOC_DIR / "TESTING.docx")


def doc_bridge():
    """Bridge_README.docx"""
    doc = Document()
    doc.add_heading("Bridge – Integration Guide", 0)
    doc.add_paragraph("Combines Breath + emotion model. Sends OSC to TouchDesigner.")
    doc.add_heading("API", level=1)
    doc.add_paragraph("prepare_audio(waveform, sample_rate) – Denoise + enhance")
    doc.add_paragraph("run_combined(audio_path=None, waveform=None, sample_rate=None, clean_audio=True)")
    doc.add_paragraph("configure_osc(ip, port, enabled)")
    doc.add_heading("OSC messages", level=1)
    doc.add_paragraph("/emotion (int 0-6), /frequency (float 0-1), /breath (int 0-3), /bpm (float 0-1)")
    doc.save(DOC_DIR / "Bridge_README.docx")


def doc_english_model():
    """English_Model_README.docx"""
    doc = Document()
    doc.add_heading("English Model", 0)
    doc.add_paragraph("emotion2vec_plus_base + noise model (DeepFilterNet2/noisereduce)")
    doc.add_heading("Modules", level=1)
    doc.add_paragraph("english_model.py – Emotion recognition")
    doc.add_paragraph("noise_model.py – Noise cleaning + enhance_audio")
    doc.add_paragraph("download_model.py – Download to model/")
    doc.add_paragraph("load_model.py – Load from model/")
    doc.add_heading("API", level=1)
    doc.add_paragraph("get_model(), get_noise_cleaner(), enhance_audio()")
    doc.save(DOC_DIR / "English_Model_README.docx")


def doc_ascended():
    """Ascended_Intelligence_README.docx"""
    doc = Document()
    doc.add_heading("ASCENDED Intelligence: Breath Detection", 0)
    doc.add_paragraph("Audio-only breath detection and F0-based emotion mapping. 100 ms chunks, 15 s history.")
    doc.add_heading("Key", level=1)
    doc.add_paragraph("BreathDetector(sample_rate=16000, use_opensmile=True)")
    doc.add_paragraph("process_chunk(audio_chunk) – Returns breath_rate_bpm, emotion (f0_hz, top_emotion)")
    doc.add_paragraph("Used by bridge for combined pipeline.")
    doc.save(DOC_DIR / "Ascended_Intelligence_README.docx")


def main():
    doc_readme()
    doc_testing()
    doc_bridge()
    doc_english_model()
    doc_ascended()
    print("Updated: README, TESTING, Bridge_README, English_Model_README, Ascended_Intelligence_README")


if __name__ == "__main__":
    main()
